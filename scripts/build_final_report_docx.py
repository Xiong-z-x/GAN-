from __future__ import annotations

import re
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
REPORT_MD = ROOT / "GAN_真实人像生成实验报告.md"
PACKAGE_DIR = ROOT / "GAN_results_images" / "final_images_package"
FINAL_DIR = ROOT / "最终报告"
FINAL_DOCX = FINAL_DIR / "GAN_真实人像生成实验报告_最终版.docx"
FINAL_PDF = FINAL_DIR / "GAN_真实人像生成实验报告_最终版.pdf"


def set_run_font(run, *, size: float | None = None, bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic


def set_paragraph_spacing(paragraph, *, before: float = 0, after: float = 6, line: float = 1.15) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def set_cell_text(cell, text: str, *, bold: bool = False, size: float = 9) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text.strip())
    set_run_font(run, size=size, bold=bold)
    set_paragraph_spacing(paragraph, after=0, line=1.0)
    for margin_name in ["top", "bottom", "left", "right"]:
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_mar = tc_pr.first_child_found_in("w:tcMar")
        if tc_mar is None:
            tc_mar = OxmlElement("w:tcMar")
            tc_pr.append(tc_mar)
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), "80")
        node.set(qn("w:type"), "dxa")


def add_cover(document: Document) -> None:
    # 封面信息与课程作业提交内容保持一致，便于最终导出后直接交付。
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("深度学习课程大作业实验报告")
    set_run_font(run, size=18, bold=True)
    set_paragraph_spacing(paragraph, before=120, after=36)

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("基于 GAN 的真实人像生成与风格迁移扩展")
    set_run_font(run, size=20, bold=True)
    set_paragraph_spacing(paragraph, after=48)

    info = [
        "姓名：熊振兴",
        "学号：2023302121197",
        "核心任务：真实人像生成",
        "基础模型：手写 DCGAN baseline",
        "复现模型：StyleGAN3 官方 FFHQ 预训练权重",
        "增强模块：AnimeGANv2 动漫化与 CycleGAN 风格迁移",
        "实验平台：AutoDL GPU 云平台",
        "报告日期：2026 年 5 月 2 日",
    ]
    for item in info:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(item)
        set_run_font(run, size=12)
        set_paragraph_spacing(paragraph, after=8)

    document.add_page_break()


def add_manual_toc(document: Document, markdown: str) -> None:
    heading = document.add_heading("目录", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for line in markdown.splitlines():
        match = re.match(r"^(#{2,3})\s+(.+)$", line)
        if not match:
            continue
        level = len(match.group(1))
        text = match.group(2).strip()
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Cm(0.75 if level == 3 else 0)
        run = paragraph.add_run(text)
        set_run_font(run, size=10.5)
        set_paragraph_spacing(paragraph, after=2, line=1.0)
    document.add_page_break()


def add_code_block(document: Document, lines: list[str]) -> None:
    for line in lines:
        paragraph = document.add_paragraph()
        run = paragraph.add_run(line)
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
        run.font.size = Pt(8.5)
        set_paragraph_spacing(paragraph, after=0, line=1.0)


def parse_markdown_table(lines: list[str]) -> list[list[str]]:
    rows: list[list[str]] = []
    for line in lines:
        stripped = line.strip()
        if re.fullmatch(r"\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?", stripped):
            continue
        cells = [cell.strip() for cell in stripped.strip("|").split("|")]
        rows.append(cells)
    return rows


def add_table(document: Document, table_lines: list[str]) -> None:
    rows = parse_markdown_table(table_lines)
    if not rows:
        return
    col_count = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"
    for r, row in enumerate(rows):
        for c in range(col_count):
            set_cell_text(table.cell(r, c), row[c] if c < len(row) else "", bold=(r == 0), size=8.5 if col_count >= 5 else 9.5)
    document.add_paragraph()


def resolve_image(path_text: str) -> Path:
    path = Path(path_text)
    if path.is_absolute():
        return path
    return ROOT / path


def add_image(document: Document, alt: str, path_text: str) -> None:
    image_path = resolve_image(path_text)
    if not image_path.exists():
        paragraph = document.add_paragraph()
        run = paragraph.add_run(f"[图片缺失] {path_text}")
        set_run_font(run, size=10, bold=True)
        return
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    max_width = Inches(6.3)
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=max_width)
    caption = document.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption.add_run(alt.strip())
    set_run_font(run, size=9.5, italic=True)
    run.font.color.rgb = RGBColor(90, 90, 90)
    set_paragraph_spacing(caption, after=10)


def add_heading(document: Document, text: str, level: int) -> None:
    heading = document.add_heading(text, level=min(level, 3))
    for run in heading.runs:
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")


def add_paragraph_text(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    set_paragraph_spacing(paragraph)
    paragraph.paragraph_format.first_line_indent = Cm(0.74)
    run = paragraph.add_run(text)
    set_run_font(run, size=10.5)


def add_list_item(document: Document, text: str, ordered: bool) -> None:
    paragraph = document.add_paragraph(style="List Number" if ordered else "List Bullet")
    run = paragraph.add_run(text)
    set_run_font(run, size=10.5)
    set_paragraph_spacing(paragraph, after=2, line=1.0)


def render_markdown(document: Document, markdown: str) -> None:
    lines = markdown.splitlines()
    i = 0
    in_code = False
    code_lines: list[str] = []
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                add_code_block(document, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        if stripped.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            add_table(document, table_lines)
            continue

        image_match = re.match(r"!\[(.*?)\]\((.*?)\)", stripped)
        if image_match:
            add_image(document, image_match.group(1), image_match.group(2))
            i += 1
            continue

        heading_match = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2)
            if level == 1:
                i += 1
                continue
            add_heading(document, text, level - 1)
            i += 1
            continue

        ordered_match = re.match(r"^\d+\.\s+(.+)$", stripped)
        if ordered_match:
            add_list_item(document, ordered_match.group(1), ordered=True)
            i += 1
            continue

        bullet_match = re.match(r"^[-*]\s+(.+)$", stripped)
        if bullet_match:
            add_list_item(document, bullet_match.group(1), ordered=False)
            i += 1
            continue

        add_paragraph_text(document, stripped)
        i += 1


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.3)
    section.right_margin = Cm(2.3)
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)


def export_pdf_with_word(docx_path: Path, pdf_path: Path) -> tuple[bool, str]:
    try:
        import win32com.client
    except Exception as exc:
        return False, f"win32com 不可用：{exc}"

    word = None
    try:
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(str(docx_path))
        doc.SaveAs(str(pdf_path), FileFormat=17)
        doc.Close(False)
        word.Quit()
        return True, "PDF 已通过 Microsoft Word 导出"
    except Exception as exc:
        if word is not None:
            try:
                word.Quit()
            except Exception:
                pass
        return False, f"PDF 导出失败：{exc}"


def main() -> None:
    FINAL_DIR.mkdir(parents=True, exist_ok=True)
    markdown = REPORT_MD.read_text(encoding="utf-8")
    document = Document()
    configure_document(document)
    add_cover(document)
    add_manual_toc(document, markdown)
    render_markdown(document, markdown)
    document.save(FINAL_DOCX)

    package_docx = PACKAGE_DIR / FINAL_DOCX.name
    package_docx.write_bytes(FINAL_DOCX.read_bytes())
    # 将正文之外的说明文件一并同步进最终包，方便后续直接提交或压缩传输。
    for source in [
        REPORT_MD,
        ROOT / "报告素材使用说明.md",
        ROOT / "最终提交清单.md",
        ROOT / "代码附录.md",
        ROOT / "文件说明.md",
        ROOT / "README.md",
        ROOT / "AutoDL_运行指南.md",
    ]:
        target = PACKAGE_DIR / source.name
        target.write_text(source.read_text(encoding="utf-8"), encoding="utf-8")

    media_count = 0
    with zipfile.ZipFile(FINAL_DOCX) as archive:
        media_count = len([name for name in archive.namelist() if name.startswith("word/media/")])

    pdf_ok, pdf_message = export_pdf_with_word(FINAL_DOCX, FINAL_PDF)
    if pdf_ok:
        package_pdf = PACKAGE_DIR / FINAL_PDF.name
        package_pdf.write_bytes(FINAL_PDF.read_bytes())

    manifest = PACKAGE_DIR / "file_manifest.txt"
    files = sorted(path.relative_to(PACKAGE_DIR).as_posix() for path in PACKAGE_DIR.rglob("*") if path.is_file())
    manifest.write_text("\n".join(files) + "\n", encoding="utf-8")

    print(f"DOCX: {FINAL_DOCX}")
    print(f"DOCX size: {FINAL_DOCX.stat().st_size}")
    print(f"embedded media: {media_count}")
    print(pdf_message)
    if FINAL_PDF.exists():
        print(f"PDF: {FINAL_PDF}")
        print(f"PDF size: {FINAL_PDF.stat().st_size}")
    print(f"package DOCX: {package_docx}")


if __name__ == "__main__":
    main()
